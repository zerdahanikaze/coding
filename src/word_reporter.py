"""
word_reporter.py — Pure Python Word report generator using python-docx.
No Node.js or npm required.

Install dependency:
    pip install python-docx
"""

import os
import tempfile
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────
PRIMARY      = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT       = RGBColor(0x2E, 0x86, 0xC1)
GREEN        = RGBColor(0x1E, 0x84, 0x49)
RED          = RGBColor(0xC0, 0x39, 0x2B)
ORANGE       = RGBColor(0xD3, 0x54, 0x00)
MID_GREY     = RGBColor(0x7F, 0x8C, 0x8D)
DARK_TEXT    = RGBColor(0x1C, 0x28, 0x33)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

BG_LIGHT_BLUE   = "D6EAF8"
BG_LIGHT_GREEN  = "D5F5E3"
BG_LIGHT_RED    = "FADBD8"
BG_LIGHT_ORANGE = "FDEBD0"
BG_GREY         = "F2F3F4"
BG_WHITE        = "FFFFFF"
BG_PRIMARY      = "1E3A5F"
BG_ACCENT       = "2E86C1"


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_row_height(row, height_twips: int):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trPr.append(trHeight)


def _para_border_bottom(para, color="2E86C1", size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Style helpers ──────────────────────────────────────────────────────────────

def _run(para, text: str, bold=False, italic=False, size_pt=11,
         color: RGBColor = None, font="Arial"):
    run = para.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _heading(doc: Document, text: str, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    _run(para, text, bold=True,
         size_pt=16 if level == 1 else 13,
         color=PRIMARY if level == 1 else ACCENT)
    if level == 1:
        _para_border_bottom(para, color="2E86C1", size=8)
    return para


def _body(doc: Document, text: str, space_before=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(4)
    _run(para, text, size_pt=10, color=DARK_TEXT)
    return para


def _spacer(doc: Document):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)


# ── Table helpers ──────────────────────────────────────────────────────────────

def _add_table(doc: Document, headers: list, rows: list,
               col_widths_in: list = None, highlight_row: int = None):
    """
    Add a styled table.
    headers       : list of column header strings
    rows          : list of lists (each inner list = one row of strings)
    col_widths_in : list of column widths in inches (must sum reasonably)
    highlight_row : 0-based index of a row to highlight green
    """
    n_cols = len(headers)
    if col_widths_in is None:
        col_widths_in = [6.5 / n_cols] * n_cols

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    _set_row_height(hdr_row, 400)
    for i, (hdr, w) in enumerate(zip(headers, col_widths_in)):
        cell = hdr_row.cells[i]
        cell.width = Inches(w)
        _set_cell_bg(cell, BG_PRIMARY)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, hdr, bold=True, size_pt=10, color=WHITE)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        _set_row_height(row, 360)
        is_highlight = (r_idx == highlight_row)
        bg = BG_LIGHT_GREEN if is_highlight else (BG_GREY if r_idx % 2 else BG_WHITE)
        for c_idx, (val, w) in enumerate(zip(row_data, col_widths_in)):
            cell = row.cells[c_idx]
            cell.width = Inches(w)
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, val, bold=is_highlight, size_pt=10,
                 color=GREEN if is_highlight else DARK_TEXT)

    return table


def _kpi_table(doc: Document, kpis: list):
    """
    kpis: list of (label, value, bg_hex) tuples, max 4 per row.
    """
    n = len(kpis)
    col_w = 6.5 / n
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    _set_row_height(row, 600)
    for i, (label, value, bg) in enumerate(kpis):
        cell = row.cells[i]
        cell.width = Inches(col_w)
        _set_cell_bg(cell, bg)
        _set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p1, value, bold=True, size_pt=16, color=PRIMARY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, label, size_pt=9, color=MID_GREY)
    return table


def _banner(doc: Document, title: str, subtitle: str, bg: str, title_color: RGBColor):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    _set_cell_bg(cell, bg)
    _set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p1, title, bold=True, size_pt=13, color=title_color)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, subtitle, size_pt=10, color=DARK_TEXT)
    return table


# ── Main public function ───────────────────────────────────────────────────────

def generate_forecast_report(
    historical_data: dict,
    forecast_data: dict,
    best_model: str,
    best_accuracy,
    model_results: dict,
    stats: dict,
    periods: int,
    selected_product: Optional[str] = None,
    product_profit_summary: Optional[list] = None,
    product_trend_summary: Optional[list] = None,
    has_product: bool = False
) -> str:
    """
    Generate a Word (.docx) forecast report using python-docx.
    Returns the path to the generated .docx file.
    """

    # ── Pre-compute summary values ─────────────────────────────────────────────
    fc_sales   = forecast_data.get('forecast_sales', [])
    fc_revenue = forecast_data.get('forecast_revenue', [])
    hist_sales   = historical_data.get('sales', [])
    hist_revenue = historical_data.get('revenue', [])

    avg_fc_sales   = sum(fc_sales)   / max(len(fc_sales), 1)
    avg_fc_revenue = sum(fc_revenue) / max(len(fc_revenue), 1)
    avg_hist_sales   = stats.get('avg_sales', sum(hist_sales) / max(len(hist_sales), 1))
    avg_hist_revenue = sum(hist_revenue) / max(len(hist_revenue), 1) if hist_revenue else 0

    sales_growth   = ((avg_fc_sales   - avg_hist_sales)   / avg_hist_sales   * 100) if avg_hist_sales   else 0
    revenue_growth = ((avg_fc_revenue - avg_hist_revenue) / avg_hist_revenue * 100) if avg_hist_revenue else 0
    accuracy_text  = f"{best_accuracy:.2f}%" if best_accuracy else "N/A"

    report_title  = "Sales & Revenue Forecast Report"
    if selected_product:
        report_title += f" — {selected_product}"
    generated_on  = datetime.now().strftime("%B %d, %Y at %H:%M")

    # ── Create document ────────────────────────────────────────────────────────
    doc = Document()

    # Page margins (1 inch all round)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Cover ──────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after  = Pt(6)
    _run(title_para, report_title, bold=True, size_pt=22, color=PRIMARY)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sub_para, "AI-Powered Sales Intelligence Report", size_pt=12, color=MID_GREY)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(date_para, f"Generated on {generated_on}", size_pt=10, italic=True, color=MID_GREY)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(6)
    divider.paragraph_format.space_after  = Pt(16)
    _para_border_bottom(divider, color="2E86C1", size=12)

    # ── 1. Executive Summary ───────────────────────────────────────────────────
    _heading(doc, "1. Executive Summary")
    scope = f'for product "{selected_product}"' if selected_product else "across all products"
    _body(doc,
        f"This report presents a {periods}-period sales and revenue forecast {scope}, "
        f"generated using the best-performing model: {best_model} (Accuracy MAPE: {accuracy_text}). "
        f"Forecasted average sales stand at ${avg_fc_sales:,.2f} and average revenue at "
        f"${avg_fc_revenue:,.2f} per period, reflecting a {sales_growth:+.1f}% growth trend "
        f"vs historical averages."
    )
    _spacer(doc)

    _kpi_table(doc, [
        ("Forecast Avg Sales",    f"${avg_fc_sales:,.0f}",    BG_LIGHT_BLUE),
        ("Forecast Avg Revenue",  f"${avg_fc_revenue:,.0f}",  BG_LIGHT_GREEN),
        ("Sales Growth",          f"{sales_growth:+.1f}%",
         BG_LIGHT_GREEN if sales_growth >= 0 else BG_LIGHT_RED),
        ("Forecast Horizon",      f"{periods} Periods",       BG_LIGHT_ORANGE),
    ])
    _spacer(doc)

    # ── 2. Historical Data Summary ─────────────────────────────────────────────
    _heading(doc, "2. Historical Data Summary")
    _body(doc,
        f"Historical dataset contains {len(hist_sales)} records. "
        f"Average sales: ${avg_hist_sales:,.2f} | "
        f"Average revenue: ${avg_hist_revenue:,.2f} | "
        f"Min sales: ${stats.get('min_sales', 0):,.2f} | "
        f"Max sales: ${stats.get('max_sales', 0):,.2f}."
    )

    # ── 3. Model Performance ───────────────────────────────────────────────────
    _heading(doc, "3. Model Performance Comparison")
    _body(doc,
        "All evaluated forecasting models are shown below ranked by accuracy (MAPE). "
        "Lower MAPE = better accuracy. The winning model is highlighted in green."
    )
    _spacer(doc)

    model_rows = []
    best_row_idx = None
    for i, (model_name, result) in enumerate(model_results.items()):
        mape = f"{result['mape']:.2f}%" if result.get('mape') else "N/A"
        rmse = f"{result['rmse']:.2f}"  if result.get('rmse') else "N/A"
        label = model_name + (" ★ Best" if model_name == best_model else "")
        model_rows.append([label, mape, rmse])
        if model_name == best_model:
            best_row_idx = i

    _add_table(doc,
        headers=["Model", "MAPE (Accuracy)", "RMSE"],
        rows=model_rows,
        col_widths_in=[3.2, 1.65, 1.65],
        highlight_row=best_row_idx
    )
    _spacer(doc)

    # ── 4. Forecast Results ────────────────────────────────────────────────────
    _heading(doc, "4. Forecast Results")
    _body(doc,
        f"Forecast generated for {periods} periods using the {best_model} model. "
        "Sales and revenue projections are shown below."
    )
    _spacer(doc)

    forecast_rows = [
        [d, f"${s:,.2f}", f"${r:,.2f}"]
        for d, s, r in zip(
            forecast_data.get('dates', []),
            fc_sales,
            fc_revenue
        )
    ]
    _add_table(doc,
        headers=["Period / Date", "Forecast Sales ($)", "Forecast Revenue ($)"],
        rows=forecast_rows,
        col_widths_in=[2.17, 2.17, 2.16]
    )
    _spacer(doc)

    # ── 5. Product Profitability (conditional) ─────────────────────────────────
    section_num = 5
    if has_product and product_profit_summary and not selected_product:
        _heading(doc, f"{section_num}. Product Profitability Analysis")
        section_num += 1

        top = product_profit_summary[0]
        top_name   = top.get('product', '')
        top_profit = top.get('Profit', 0)
        top_margin = top.get('Profit_Margin_%', 0)

        _body(doc,
            f'Across all products, "{top_name}" leads in profitability with a total profit of '
            f"${top_profit:,.2f} and a margin of {top_margin:.2f}%. "
            "Products are ranked below from most to least profitable."
        )
        _spacer(doc)

        # Top product banner
        _banner(doc,
            title=f"🏆 Top Performing Product: {top_name}",
            subtitle=f"Profit: ${top_profit:,.2f}   |   Margin: {top_margin:.2f}%",
            bg=BG_LIGHT_GREEN,
            title_color=GREEN
        )
        _spacer(doc)

        profit_rows = []
        best_profit_idx = 0
        for i, p in enumerate(product_profit_summary):
            profit_rows.append([
                p.get('product', ''),
                f"${p.get('Total_Sales', 0):,.2f}",
                f"${p.get('Total_Revenue', 0):,.2f}",
                f"${p.get('Profit', 0):,.2f}",
                f"{p.get('Profit_Margin_%', 0):.2f}%",
            ])

        _add_table(doc,
            headers=["Product", "Total Sales", "Total Revenue", "Profit", "Margin %"],
            rows=profit_rows,
            col_widths_in=[1.8, 1.2, 1.3, 1.1, 1.1],
            highlight_row=best_profit_idx
        )
        _spacer(doc)

    # ── 6. Product Sales Trend Review (conditional) ───────────────────────────
    if has_product and product_trend_summary and not selected_product:
        _heading(doc, f"{section_num}. Product Sales Trend Review")
        section_num += 1

        _body(doc,
            "This section reviews the historical sales trend for each product by comparing "
            "the first half vs the second half of the available data period. Growth rate is "
            "the percentage change in average sales between the two halves."
        )
        _spacer(doc)

        trend_rows = []
        for t in product_trend_summary:
            trend_rows.append([
                t.get('product', ''),
                f"${t.get('avg_sales', 0):,.2f}",
                f"${t.get('avg_revenue', 0):,.2f}",
                f"{t.get('growth_pct', 0):+.2f}%",
                t.get('trend', ''),
            ])

        _add_table(doc,
            headers=["Product", "Avg Sales", "Avg Revenue", "Growth Rate", "Trend Signal"],
            rows=trend_rows,
            col_widths_in=[1.8, 1.2, 1.3, 1.1, 1.1]
        )

        legend = doc.add_paragraph()
        legend.paragraph_format.space_before = Pt(6)
        _run(legend,
            "Legend:  📈 Growing = >+5%   |   📉 Declining = >−5%   |   ➡️ Stable = within ±5%",
            italic=True, size_pt=9, color=MID_GREY
        )
        _spacer(doc)

    # ── 7. Recommendations ────────────────────────────────────────────────────
    _heading(doc, f"{section_num}. Recommendations")

    rec_items = [
        f"Model Selection: Continue using {best_model} for future forecasts "
        f"given its superior accuracy (MAPE: {accuracy_text}).",
        f"Forecast Horizon: The current {periods}-period forecast is suitable for near-term "
        "planning. For longer-term strategy, consider extending to 12–24 periods.",
        "Data Quality: Ensure historical data is consistently updated to maintain forecast "
        "accuracy over time.",
        "Review Cycle: Re-run this forecast monthly or after any major business event to keep "
        "projections current.",
    ]
    if has_product and product_profit_summary:
        top_name = product_profit_summary[0].get('product', '')
        rec_items.insert(2,
            f'Product Focus: Prioritise "{top_name}" in marketing and inventory planning — '
            "it delivers the highest profitability."
        )
        rec_items.insert(3,
            "Declining Products: Investigate products flagged as Declining in the Trend "
            "Review and consider promotional actions or discontinuation."
        )

    for item in rec_items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after  = Pt(3)
        _run(para, item, size_pt=10, color=DARK_TEXT)

    # ── Footer note ────────────────────────────────────────────────────────────
    _spacer(doc)
    footer_para = doc.add_paragraph()
    footer_para.paragraph_format.space_before = Pt(16)
    _para_border_bottom(footer_para, color="BDC3C7", size=4)
    _run(footer_para,
        "This report was automatically generated by the Smart Sales Forecasting Dashboard. "
        "Forecasts are statistical estimates and should be used as a guide alongside "
        "business judgment.",
        italic=True, size_pt=9, color=MID_GREY
    )

    # ── Save ───────────────────────────────────────────────────────────────────
    output_dir  = tempfile.mkdtemp()
    output_path = os.path.join(
        output_dir,
        f"forecast_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    doc.save(output_path)
    return output_path